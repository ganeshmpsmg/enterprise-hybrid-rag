"""
Document Loader - Unified multi-format document loading interface.
Supports PDF, DOCX, TXT, Markdown with automatic format detection.
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from src.ingestion.pdf_loader import PDFLoader, PDFDocument

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """
    Unified document representation regardless of source format.
    This is the standard document object passed through the pipeline.
    """

    doc_id: str  # Unique document identifier (hash)
    source: str  # Original file path
    file_name: str  # File name
    file_type: str  # pdf | docx | txt | md
    content: str  # Full extracted text content
    metadata: dict  # Document-level metadata
    page_contents: list[dict] = field(default_factory=list)  # Per-page content
    total_pages: int = 1
    word_count: int = 0
    char_count: int = 0

    def __post_init__(self):
        if self.word_count == 0:
            self.word_count = len(self.content.split())
        if self.char_count == 0:
            self.char_count = len(self.content)


class DocumentLoader:
    """
    Unified document loader that routes to format-specific loaders.

    Supported formats:
    - PDF (.pdf)      -> PDFLoader
    - Word  (.docx)   -> DocxLoader
    - Text  (.txt)    -> TextLoader
    - Markdown (.md)  -> MarkdownLoader

    Usage:
        loader = DocumentLoader()
        doc = loader.load("research_paper.pdf")
        docs = loader.load_directory("./papers/", recursive=True)
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}

    def __init__(
        self,
        extract_tables: bool = True,
        encoding: str = "utf-8",
        pdf_password: Optional[str] = None,
    ):
        self.extract_tables = extract_tables
        self.encoding = encoding
        self.pdf_password = pdf_password
        self._pdf_loader = PDFLoader(
            extract_tables=extract_tables,
            password=pdf_password,
        )

    def load(self, file_path: str | Path) -> Document:
        """
        Load a single document file.

        Args:
            file_path: Path to the document

        Returns:
            Unified Document object

        Raises:
            ValueError: For unsupported file types
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        logger.info(f"Loading document: {path.name} (type: {ext})")

        if ext == ".pdf":
            return self._load_pdf(path)
        elif ext in {".docx", ".doc"}:
            return self._load_docx(path)
        elif ext == ".txt":
            return self._load_text(path)
        elif ext in {".md", ".markdown"}:
            return self._load_markdown(path)
        else:
            return self._load_text(path)  # Generic fallback

    def _load_pdf(self, path: Path) -> Document:
        """Load PDF via PDFLoader and convert to unified Document."""
        pdf_doc: PDFDocument = self._pdf_loader.load(path)
        page_contents = [
            {
                "page_number": p.page_number,
                "content": p.text,
                "word_count": p.word_count,
                "has_tables": p.has_tables,
            }
            for p in pdf_doc.pages
        ]
        return Document(
            doc_id=pdf_doc.doc_hash,
            source=pdf_doc.file_path,
            file_name=pdf_doc.file_name,
            file_type="pdf",
            content=pdf_doc.raw_text,
            metadata={**pdf_doc.metadata, "file_type": "pdf"},
            page_contents=page_contents,
            total_pages=pdf_doc.total_pages,
        )

    def _load_docx(self, path: Path) -> Document:
        """Load DOCX using python-docx."""
        import hashlib

        try:
            from docx import Document as DocxDocument

            docx = DocxDocument(str(path))
            paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            # Extract tables
            for table in docx.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content += f"\n{row_text}"

            doc_hash = hashlib.sha256(path.read_bytes()).hexdigest()
            core_props = docx.core_properties
            metadata = {
                "source": str(path.absolute()),
                "file_name": path.name,
                "file_type": "docx",
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "file_size_bytes": path.stat().st_size,
            }
            return Document(
                doc_id=doc_hash,
                source=str(path.absolute()),
                file_name=path.name,
                file_type="docx",
                content=content,
                metadata=metadata,
                total_pages=1,
            )
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise

    def _load_text(self, path: Path) -> Document:
        """Load plain text file."""
        import hashlib

        content = path.read_text(encoding=self.encoding, errors="replace")
        doc_hash = hashlib.sha256(path.read_bytes()).hexdigest()
        return Document(
            doc_id=doc_hash,
            source=str(path.absolute()),
            file_name=path.name,
            file_type="txt",
            content=content,
            metadata={
                "source": str(path.absolute()),
                "file_name": path.name,
                "file_type": "txt",
                "file_size_bytes": path.stat().st_size,
            },
        )

    def _load_markdown(self, path: Path) -> Document:
        """Load Markdown file, stripping markdown syntax."""
        import hashlib
        import re

        raw = path.read_text(encoding=self.encoding, errors="replace")
        # Strip basic markdown: headers, bold, italic, code blocks, links
        content = re.sub(r"#{1,6}\s+", "", raw)  # Remove headers
        content = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", content)  # Bold/italic
        content = re.sub(r"`{3}[^`]*`{3}", "", content, flags=re.DOTALL)  # Code blocks
        content = re.sub(r"`([^`]+)`", r"\1", content)  # Inline code
        content = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", content)  # Links
        doc_hash = hashlib.sha256(path.read_bytes()).hexdigest()
        return Document(
            doc_id=doc_hash,
            source=str(path.absolute()),
            file_name=path.name,
            file_type="md",
            content=content,
            metadata={
                "source": str(path.absolute()),
                "file_name": path.name,
                "file_type": "md",
                "file_size_bytes": path.stat().st_size,
                "raw_markdown": raw,
            },
        )

    def load_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
        glob_pattern: Optional[str] = None,
    ) -> list[Document]:
        """
        Load all supported documents from a directory.

        Args:
            directory: Directory to scan
            recursive: If True, scan subdirectories
            glob_pattern: Optional custom glob pattern

        Returns:
            List of loaded Document objects
        """
        dir_path = Path(directory)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {directory}")

        if glob_pattern:
            paths = list(
                dir_path.glob(glob_pattern)
                if not recursive
                else dir_path.rglob(glob_pattern)
            )
        else:
            paths = []
            pattern = "**/*" if recursive else "*"
            for ext in self.SUPPORTED_EXTENSIONS:
                paths.extend(dir_path.glob(f"{pattern}{ext}"))

        logger.info(f"Found {len(paths)} files in {directory}")
        documents = []
        for p in sorted(paths):
            try:
                doc = self.load(p)
                documents.append(doc)
                logger.debug(f"Loaded: {p.name}")
            except Exception as e:
                logger.error(f"Failed to load {p.name}: {e}")
        logger.info(f"Successfully loaded {len(documents)}/{len(paths)} documents")
        return documents
